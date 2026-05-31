#!/usr/bin/env python3
"""
Duo Parser: Extract and organize MCQ data from .docx files
Converts Word documents into a structured JSON database optimized for the Duolingo-style learning app.

Usage:
    python scripts/duo_parser.py

Expected folder structure:
    subjects_vault/
    ├── Subject1/
    │   ├── file1.docx
    │   └── file2.docx
    ├── Subject2/
    │   └── file1.docx
"""

import os
import json
from pathlib import Path
from docx import Document
from typing import Dict, List, Any


class MCQParser:
    """Parse MCQ data from .docx files in subjects_vault directory."""
    
    QUESTIONS_PER_LEVEL = 10
    VAULT_PATH = Path("subjects_vault")
    OUTPUT_PATH = Path("src/data/db.json")
    
    def __init__(self):
        self.subjects: Dict[str, List[List[Dict[str, Any]]]] = {}
    
    def parse_all_subjects(self) -> None:
        """Parse all .docx files from subjects_vault directory."""
        if not self.VAULT_PATH.exists():
            print(f"❌ Error: {self.VAULT_PATH} directory not found!")
            return
        
        subject_folders = [d for d in self.VAULT_PATH.iterdir() if d.is_dir()]
        
        if not subject_folders:
            print("⚠️  Warning: No subject folders found in subjects_vault/")
            return
        
        print(f"🔍 Found {len(subject_folders)} subject(s)")
        
        for subject_folder in sorted(subject_folders):
            subject_name = subject_folder.name
            print(f"\n📚 Processing subject: {subject_name}")
            
            docx_files = list(subject_folder.glob("*.docx"))
            if not docx_files:
                print(f"   ⚠️  No .docx files found in {subject_name}/")
                continue
            
            questions = []
            for docx_file in docx_files:
                print(f"   📄 Reading: {docx_file.name}")
                extracted = self.extract_from_docx(docx_file)
                questions.extend(extracted)
                print(f"      ✓ Extracted {len(extracted)} questions")
            
            if questions:
                levels = self.chunk_into_levels(questions)
                self.subjects[subject_name] = levels
                print(f"   ✅ Created {len(levels)} level(s) with {len(questions)} total questions")
    
    def extract_from_docx(self, docx_path: Path) -> List[Dict[str, Any]]:
        """Extract MCQ data from a .docx file."""
        doc = Document(docx_path)
        questions = []
        
        i = 0
        while i < len(doc.paragraphs):
            para = doc.paragraphs[i]
            text = para.text.strip()
            
            # Look for question patterns (typically starting with "Q:", "Question:", or a number)
            if self._is_question_start(text) and i + 4 < len(doc.paragraphs):
                question_data = self._extract_single_question(doc.paragraphs, i)
                if question_data:
                    questions.append(question_data)
                    i += 6  # Skip past the extracted question
                else:
                    i += 1
            else:
                i += 1
        
        return questions
    
    def _is_question_start(self, text: str) -> bool:
        """Check if text looks like the start of a question."""
        return (
            text.lower().startswith(("q:", "question:", "q."))
            or (len(text) > 0 and text[0].isdigit() and len(text) < 100)
        )
    
    def _extract_single_question(
        self, paragraphs: List, start_idx: int
    ) -> Dict[str, Any] | None:
        """Extract a single question from paragraph sequence."""
        try:
            question_text = paragraphs[start_idx].text.strip()
            
            # Extract the actual question (remove "Q:" or number prefix)
            question_text = self._clean_question_text(question_text)
            
            # Assuming next 4 lines are options A, B, C, D
            options = []
            correct_answer = None
            
            for j in range(1, 5):
                if start_idx + j >= len(paragraphs):
                    return None
                
                option_text = paragraphs[start_idx + j].text.strip()
                
                # Parse option format: "A) option text" or "A: option text"
                if len(option_text) > 2 and option_text[0] in "ABCD" and option_text[1] in "):":
                    letter = option_text[0]
                    option_content = option_text[2:].strip()
                    options.append({"letter": letter, "text": option_content})
                    
                    # Check if this is marked as correct (e.g., asterisk, bold, etc.)
                    if "*" in option_text or "✓" in option_text or "correct" in option_text.lower():
                        correct_answer = letter
                else:
                    return None
            
            # If no explicit correct answer marked, assume first option (placeholder)
            if not correct_answer and options:
                correct_answer = options[0]["letter"]
            
            # Extract explanation (usually after options)
            explanation = ""
            if start_idx + 5 < len(paragraphs):
                explanation = paragraphs[start_idx + 5].text.strip()
            
            return {
                "id": f"q_{hash(question_text) % 1000000}",
                "question": {
                    "en": question_text,
                    "ar": question_text,  # Placeholder - should be translated
                },
                "options": options,
                "correctAnswer": correct_answer,
                "explanation": {
                    "en": explanation,
                    "ar": explanation,  # Placeholder - should be translated
                },
            }
        except Exception as e:
            print(f"      ⚠️  Error extracting question: {e}")
            return None
    
    def _clean_question_text(self, text: str) -> str:
        """Remove question number/label prefix."""
        if text.lower().startswith("q"):
            text = text[1:].lstrip(":. ")
        elif text[0].isdigit():
            text = text.lstrip("0123456789").lstrip(":.) ")
        return text
    
    def chunk_into_levels(self, questions: List[Dict[str, Any]]) -> List[List[Dict[str, Any]]]:
        """Chunk questions into levels (QUESTIONS_PER_LEVEL per level)."""
        levels = []
        for i in range(0, len(questions), self.QUESTIONS_PER_LEVEL):
            level = questions[i : i + self.QUESTIONS_PER_LEVEL]
            levels.append(level)
        return levels
    
    def save_to_json(self) -> None:
        """Save parsed data to db.json."""
        self.OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
        
        # Structure: { subjects: { "Subject Name": [ [questions in level 1], [questions in level 2] ] } }
        output_data = {
            "subjects": self.subjects,
            "metadata": {
                "totalSubjects": len(self.subjects),
                "totalQuestions": sum(
                    len(q) for subject in self.subjects.values() for q in subject
                ),
                "questionsPerLevel": self.QUESTIONS_PER_LEVEL,
            },
        }
        
        with open(self.OUTPUT_PATH, "w", encoding="utf-8") as f:
            json.dump(output_data, f, indent=2, ensure_ascii=False)
        
        print(f"\n✅ Data successfully saved to {self.OUTPUT_PATH}")
        print(f"   📊 Total subjects: {output_data['metadata']['totalSubjects']}")
        print(f"   📊 Total questions: {output_data['metadata']['totalQuestions']}")


def main():
    """Main entry point."""
    print("🚀 MCQ Parser - Duolingo Style Data Pipeline\n")
    
    parser = MCQParser()
    parser.parse_all_subjects()
    parser.save_to_json()
    
    print("\n✨ Parse complete!")


if __name__ == "__main__":
    main()
